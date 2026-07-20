from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DOCX = Path(__file__).with_name("RESPONSE_LETTER_DRAFT.docx")


def set_cell_text(cell, text: str, *, bold: bool, size: float) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)


def main() -> None:
    doc = Document(DOCX)
    table = doc.tables[0]
    header = [cell.text.strip() for cell in table.rows[0].cells[:3]]
    if header != ["Group", "Model", "Features"]:
        raise ValueError(f"Unexpected Table 1 header: {header}")

    for row_index, row in enumerate(table.rows):
        values = [cell.text.strip() for cell in row.cells[:3]]
        merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
        if row_index == 0:
            set_cell_text(merged, "Group-Model-Features", bold=True, size=7.3)
        else:
            set_cell_text(merged, "-".join(values), bold=False, size=7.3)

    doc.save(DOCX)
    print(f"Updated {DOCX}")


if __name__ == "__main__":
    main()
