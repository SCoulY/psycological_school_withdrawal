from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reviewer_response_experiments" / "RESPONSE_LETTER_DRAFT.md"
OUTPUT = ROOT / "reviewer_response_experiments" / "RESPONSE_LETTER_DRAFT.docx"
FIG = ROOT / "reviewer_response_experiments" / "figures"
LATEX_FIG = ROOT / "reviewer_response_experiments" / "latex_revision" / "source" / "figs"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:" + key), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clean_math(text: str) -> str:
    """Make the small amount of TeX in the Markdown readable in Word."""
    text = text.replace(r"\mathrm{reentry}", "reentry")
    text = text.replace(r"\mathrm{withdrawal}", "withdrawal")
    text = text.replace(r"\mathrm{W}", "W").replace(r"\mathrm{R}", "R")
    text = text.replace(r"\leq", "≤").replace(r"\geq", "≥")
    text = text.replace(r"\times", "×").replace(r"\sum", "Σ")
    text = text.replace(r"\log", "log")
    text = text.replace("\\", "")
    text = re.sub(r"([A-Za-zΣ])_\{([^{}]+)\}", r"\1_\2", text)
    text = text.replace("{", "").replace("}", "")
    return text


def add_inline_runs(paragraph, text: str) -> None:
    """Parse the limited Markdown inline syntax used by the response letter."""
    text = clean_math(text)
    tick = chr(96)
    pattern = re.compile(
        r"(\*\*.*?\*\*|\*.*?\*|"
        + re.escape(tick)
        + r".*?"
        + re.escape(tick)
        + r"|\\\(.*?\\\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith(tick):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(70, 70, 70)
        else:
            run = paragraph.add_run(token[2:-2])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc: Document, text: str, style: str = "Normal", indent: bool = False):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
    add_inline_runs(p, text)
    return p


def parse_table(lines, start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


TABLE_CAPTIONS = {
    2: "Independent outer-fold withdrawal-status AUROC for KDE aggregation rules and reference scores across age groups.",
}


def add_data_table(doc: Document, rows, table_number: int):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(clean_math(value))
            run.font.name = "Aptos"
            run.font.size = Pt(7.3 if len(rows[0]) >= 9 else 8)
            if r_idx == 0:
                run.bold = True
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    title = TABLE_CAPTIONS.get(
        table_number,
        "Quantitative results reported in the response letter.",
    )
    cap_run = caption.add_run(f"Table {table_number}. {title}")
    cap_run.italic = True
    cap_run.font.size = Pt(8.5)
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(80, 80, 80)
    return p


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.25):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_calibration_strip(doc: Document):
    path = LATEX_FIG / "calibration_combined.png"
    add_figure(
        doc,
        path,
        "Figure 2. Combined calibration curves for the revised age-stratified models (adults, teens, and children).",
        6.55,
    )


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Response letter  |  ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(100, 100, 100)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    for style_name in ("Normal", "Body Text"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.08
        style.paragraph_format.space_after = Pt(6)
    for style_name, size, color in (("Title", 16, "17365D"), ("Heading 1", 13, "17365D"), ("Heading 2", 11.5, "1F4E79")):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
    add_footer(section)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_inline_runs(p, text)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Point-to-point response letter"
    doc.core_properties.subject = "Response to Reviewer 1 and Reviewer 2"
    doc.core_properties.author = "The Authors"

    table_count = 0
    current_comment = ""
    inserted_comment1_figures = False
    inserted_comment5_figures = False
    inserted_comment7_figures = False
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, stripped[2:].strip())
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = subtitle.add_run("Point-to-point response to the Editor and Reviewers")
            rr.italic = True
            rr.font.size = Pt(10)
            rr.font.color.rgb = RGBColor(90, 90, 90)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            current_comment = ""
            i += 1
            continue
        if stripped.startswith("### "):
            if "shap radar" in current_comment and not inserted_comment7_figures:
                add_figure(doc, FIG / "p2_standard_shap" / "Fig6_standard_shap_beeswarm.png", "Figure 5. Standard SHAP beeswarm plots for the descriptive age-stratified refits.", 5.75)
                add_figure(doc, FIG / "p2_standard_shap" / "FigS_standard_shap_bar.png", "Figure 6. Mean-absolute SHAP bar plots supplied as a supplementary readability check.", 5.75)
                inserted_comment7_figures = True
            heading = stripped[4:].strip()
            add_heading(doc, heading, 2)
            current_comment = heading.lower()
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                table_count += 1
                add_data_table(doc, rows, table_count)
                if ("surpassed" in current_comment or "claim that reduced models" in current_comment) and not inserted_comment1_figures:
                    add_figure(doc, FIG / "p0_comparison" / "paired_full_vs_top10_forest.png", "Figure 1. Paired participant-level bootstrap AUROC differences between the full and top-10 feature sets.", 6.25)
                    add_calibration_strip(doc)
                    inserted_comment1_figures = True
                elif "quantitative validation of kde" in current_comment and not inserted_comment5_figures:
                    add_figure(doc, LATEX_FIG / "KDE_aggregation_AUROC.png", "Figure 3. Independent outer-fold validation of alternative KDE aggregation scores.", 6.25)
                    add_figure(doc, LATEX_FIG / "KDE_sensitivity_AUROC.png", "Figure 4. KDE score sensitivity to threshold, minimum reference-cohort size, and bandwidth.", 6.25)
                    inserted_comment5_figures = True
            continue
        if stripped.startswith("**Changes.**"):
            if "shap radar" in current_comment and not inserted_comment7_figures:
                add_figure(doc, FIG / "p2_standard_shap" / "Fig6_standard_shap_beeswarm.png", "Figure 5. Standard SHAP beeswarm plots for the descriptive age-stratified refits.", 5.75)
                add_figure(doc, FIG / "p2_standard_shap" / "FigS_standard_shap_bar.png", "Figure 6. Mean-absolute SHAP bar plots supplied as a supplementary readability check.", 5.75)
                inserted_comment7_figures = True
            p = doc.add_paragraph(style="Normal")
            r = p.add_run("Changes. ")
            r.bold = True
            add_inline_runs(p, stripped[len("**Changes.**"):].strip())
            i += 1
            continue
        if stripped == "Sincerely,":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.add_run("Sincerely,")
            i += 1
            continue
        if stripped == "The Authors":
            doc.add_paragraph("The Authors")
            i += 1
            continue
        paragraph_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt or nxt.startswith(("# ", "## ", "### ", "|", "**Changes.**")):
                break
            paragraph_lines.append(nxt)
            j += 1
        add_paragraph(doc, " ".join(paragraph_lines), indent=(current_comment != ""))
        i = j

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
